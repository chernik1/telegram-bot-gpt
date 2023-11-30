from settings import *
from docx import Document


DATE_ROW = []


def data_get() -> Document:
    """Функция получает данные из docx-файла"""
    doc = Document(FILE_NAME)
    tables = doc.tables
    table = tables[INDEX_TABLE]
    return table

def change_row(row: list[str]) -> list[str]:
    """Функция изменяет строку и формирует новую"""
    print(row[1], row[2])
    print(*DATE_ROW)

    list_presence = list(map(str, input().split()))

    if list_presence[0] in Commands.COMMANDS:
        return list_presence[0]

    new_list_presence = []
    for sign in list_presence:
        if sign == '-':
            new_list_presence.append(' ')
            new_list_presence.append(' ')
        elif sign == '++':
            new_list_presence.append('+')
            new_list_presence.append('+')
        elif sign == '--':
            new_list_presence.append(' ')
            new_list_presence.append(' ')
        elif sign == '+':
            new_list_presence.append('+')
            new_list_presence.append(' ')
        elif sign == '+-':
            new_list_presence.append('+')
            new_list_presence.append(' ')
        else:
            new_list_presence.append(sign)
            new_list_presence.append(' ')
    new_row = []
    for index in range(INDEX_NAME + 1):
        new_row.append(row[index])
    new_row.extend(new_list_presence)
    new_row.append(row[-1])
    if len(new_row) != len(row):
        raise SyntaxError('Число столбцов не совпадает. Возможно проблема в неправильном заполнении таблицы')
    return new_row


def data_form(table: Document) -> list[list]:
    """Функция формирует таблицу из docx-файла"""
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        data.append(row_data)
    return data


def data_change(data: list[list], func) -> list[list]:
    """Функция изменяет данные и формирует новую таблицу"""

    def data_check_name(name: str) -> bool:
        """Функция проверяет наличие имени"""
        if '’' in name:
            return True
        if name == name.title() and len(name.split()) == INDEX_NAME:
            return True
        return False

    def date_row_change(date_data: list) -> list:
        """Функция изменяет строку с датой"""
        new_date_data = []
        for date in date_data:
            if date not in new_date_data and date.isdigit():
                new_date_data.append(date)
        return new_date_data

    global DATE_ROW
    DATE_ROW = data[LINE_TABLE_DATE][COLUMN_TABLE_DATE:]
    DATE_ROW = date_row_change(DATE_ROW)

    new_data = []

    for row in data:
        if data_check_name(row[INDEX_NAME]):
            new_row = func(row)

            if new_row in Commands.COMMANDS:
                if new_row == 'exit':
                    raise SyntaxError('exit')
                elif new_row == 'save':
                    len_data_check = len(new_data)
                    while len_data_check != len(data):
                        new_data.append(data[len_data_check])
                        len_data_check += 1
                    flag = data_write(new_data)
                    if flag:
                        raise SyntaxError('Сохранилось')
                elif new_row == 'change':
                    index = int(input('С какой строки продолжить?(включительно) '))
                    for row_change in data[index:]:
                        if data_check_name(row_change[INDEX_NAME]):
                            new_row = func(row_change)
                            new_data.append(new_row)
                        else:
                            new_data.append(row_change)

                    if len(new_data) != len(data):
                        raise SyntaxError(
                            'Число столбцов не совпадает. Возможно проблема в неправильном заполнении таблицы')

                    return new_data

            new_data.append(new_row)
        else:
            new_data.append(row)

    if len(new_data) != len(data):
        raise SyntaxError('Число столбцов не совпадает. Возможно проблема в неправильном заполнении таблицы')

    return new_data


def data_write(new_data: list[list]):
    doc = Document(FILE_NAME)

    table = doc.tables[0]

    cell_styles = []
    for row in table.rows:
        for cell in row.cells:
            cell_styles.append(cell.paragraphs[0].style)

    for i in range(len(table.rows) - 1, 0, -1):
        table._tbl.remove(table.rows[i]._tr)

    for row_data, style in zip(new_data[1:], cell_styles):
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(cell_data)
            cell.paragraphs[0].style = style

    doc.save(FILE_SAVE)
    return True


def main():
    """Запуск и логика"""
    table = data_get()
    data = data_form(table)
    # в дате лежит вся таблица
    new_data = data_change(data, change_row)
    flag = data_write(new_data)
    return flag
